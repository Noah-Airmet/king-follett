#!/usr/bin/env python3
"""
sync_html_to_docx.py — Append sections VII, VIII, IX from the HTML edition
to the existing DOCX (which contains sections I–VI).

Uses python-docx and lxml. Preserves existing formatting conventions:
  - Heading 1 for major sections (h2)
  - Heading 2 for subsections (h3)
  - Normal style (Georgia 11pt) for body text
  - Tables with header row
  - Hyperlinks via OxmlElement
"""

import copy
import html
import os
import re
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.shared import Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree


# ── Paths ────────────────────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(SCRIPT_DIR, "king-follett-critical-edition.docx")
HTML_PATH = os.path.join(SCRIPT_DIR, "..", "docs", "index.html")
BACKUP_PATH = DOCX_PATH + ".bak"


# ── HTML parsing helpers ─────────────────────────────────────────────────

def load_html():
    """Load and parse the HTML file."""
    with open(HTML_PATH, "r", encoding="utf-8") as f:
        content = f.read()
    # Extract just the <main> content to avoid parsing JS/CSS issues
    # Use lxml HTML parser which is more forgiving
    parser = etree.HTMLParser(encoding="utf-8")
    tree = etree.fromstring(content.encode("utf-8"), parser)
    return tree


def find_sections_vii_to_ix(tree):
    """Return list of lxml elements from VII heading through end of IX."""
    # Find all h2 elements
    h2s = tree.xpath("//h2")
    start_el = None
    for h2 in h2s:
        text = (h2.text or "") + "".join(
            (c.text or "") + (c.tail or "") for c in h2
        )
        if "VII." in text or "apparatus-criticus" in (h2.get("id") or ""):
            start_el = h2
            break

    if start_el is None:
        raise RuntimeError("Could not find section VII in HTML")

    # Collect all elements from VII heading to end of main
    main = start_el.getparent()
    elements = []
    collecting = False
    for el in main:
        if el is start_el:
            collecting = True
        if collecting:
            # Stop at the final colophon (centered paragraph)
            if el.tag == "p":
                style = el.get("style", "")
                if "text-align:center" in style and "color:#888" in style:
                    break
            # Skip <hr> dividers
            if el.tag == "hr":
                continue
            elements.append(el)
    return elements


# ── DOCX formatting helpers ──────────────────────────────────────────────

def set_run_font(run, name="Georgia", size=Pt(11), bold=None, italic=None,
                 color=None):
    """Set font properties on a run."""
    run.font.name = name
    run.font.size = size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_hyperlink(paragraph, url, text, font_name="Georgia",
                  font_size=Pt(11)):
    """Add a clickable hyperlink to a paragraph using OxmlElement."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Font name
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rPr.append(rFonts)

    # Font size
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_size.pt * 2)))  # half-points
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(font_size.pt * 2)))
    rPr.append(szCs)

    # Blue color + underline (standard hyperlink style)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "0563C1")
    rPr.append(color_el)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def decode_entities(text):
    """Decode HTML entities to Unicode characters."""
    if text is None:
        return ""
    return html.unescape(text)


def get_element_text(el):
    """Get all text from an element (including children), decoded."""
    return decode_entities("".join(el.itertext()))


# ── Inline content rendering ────────────────────────────────────────────

def render_inline(paragraph, el, font_size=Pt(11)):
    """
    Recursively render inline HTML content into a Word paragraph.
    Handles <b>, <i>, <a>, and plain text.
    """
    # Process element's own text
    if el.text:
        text = decode_entities(el.text)
        if text:
            _add_text_run(paragraph, text, font_size,
                          bold=_in_tag(el, "b"),
                          italic=_in_tag(el, "i"))

    for child in el:
        tag = child.tag
        if tag == "a":
            href = child.get("href", "")
            link_text = get_element_text(child)
            if href and link_text:
                add_hyperlink(paragraph, href, link_text,
                              font_size=font_size)
        elif tag in ("b", "strong"):
            _render_bold_or_italic(paragraph, child, font_size, bold=True)
        elif tag in ("i", "em"):
            _render_bold_or_italic(paragraph, child, font_size, italic=True)
        elif tag == "br":
            run = paragraph.add_run("\n")
            set_run_font(run, size=font_size)
        else:
            # Recurse for unknown inline tags
            render_inline(paragraph, child, font_size)

        # Process tail text (text after closing tag)
        if child.tail:
            text = decode_entities(child.tail)
            if text:
                _add_text_run(paragraph, text, font_size,
                              bold=_in_tag(el, "b"),
                              italic=_in_tag(el, "i"))


def _in_tag(el, tag_name):
    """Check if el itself is the given tag."""
    return el.tag == tag_name


def _add_text_run(paragraph, text, font_size, bold=False, italic=False):
    """Add a plain text run with formatting."""
    run = paragraph.add_run(text)
    set_run_font(run, size=font_size, bold=bold if bold else None,
                 italic=italic if italic else None)


def _render_bold_or_italic(paragraph, el, font_size, bold=False, italic=False):
    """Render a <b> or <i> element and its children."""
    if el.text:
        text = decode_entities(el.text)
        if text:
            run = paragraph.add_run(text)
            set_run_font(run, size=font_size,
                         bold=bold or None,
                         italic=italic or None)

    for child in el:
        tag = child.tag
        if tag == "a":
            href = child.get("href", "")
            link_text = get_element_text(child)
            if href and link_text:
                add_hyperlink(paragraph, href, link_text,
                              font_size=font_size)
        elif tag in ("i", "em"):
            # bold+italic
            if child.text:
                text = decode_entities(child.text)
                if text:
                    run = paragraph.add_run(text)
                    set_run_font(run, size=font_size,
                                 bold=bold or None, italic=True)
            # Handle children inside <i> within <b>
            for grandchild in child:
                _render_bold_or_italic(paragraph, grandchild, font_size,
                                       bold=bold, italic=True)
                if grandchild.tail:
                    text = decode_entities(grandchild.tail)
                    if text:
                        run = paragraph.add_run(text)
                        set_run_font(run, size=font_size,
                                     bold=bold or None, italic=True)
            if child.tail:
                text = decode_entities(child.tail)
                if text:
                    run = paragraph.add_run(text)
                    set_run_font(run, size=font_size,
                                 bold=bold or None, italic=italic or None)
        elif tag in ("b", "strong"):
            # italic+bold
            if child.text:
                text = decode_entities(child.text)
                if text:
                    run = paragraph.add_run(text)
                    set_run_font(run, size=font_size,
                                 bold=True, italic=italic or None)
            for grandchild in child:
                _render_bold_or_italic(paragraph, grandchild, font_size,
                                       bold=True, italic=italic)
                if grandchild.tail:
                    text = decode_entities(grandchild.tail)
                    if text:
                        run = paragraph.add_run(text)
                        set_run_font(run, size=font_size,
                                     bold=True, italic=italic or None)
            if child.tail:
                text = decode_entities(child.tail)
                if text:
                    run = paragraph.add_run(text)
                    set_run_font(run, size=font_size,
                                 bold=bold or None, italic=italic or None)
        else:
            render_inline(paragraph, child, font_size)
            if child.tail:
                text = decode_entities(child.tail)
                if text:
                    run = paragraph.add_run(text)
                    set_run_font(run, size=font_size,
                                 bold=bold or None, italic=italic or None)


# ── Section rendering ────────────────────────────────────────────────────

def add_heading_1(doc, text):
    """Add a Heading 1 paragraph (for h2 / major sections)."""
    p = doc.add_heading(text, level=1)
    return p


def add_heading_2(doc, text):
    """Add a Heading 2 paragraph (for h3 / subsections)."""
    p = doc.add_heading(text, level=2)
    return p


def add_body_paragraph(doc, el, font_size=Pt(11)):
    """Add a Normal-style paragraph with inline formatting from HTML element."""
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    render_inline(p, el, font_size)
    return p


def add_table(doc, table_el):
    """Convert an HTML <table> into a Word table."""
    rows_data = []
    header_row = None

    for tr in table_el.iter("tr"):
        cells = []
        is_header = False
        for cell in tr:
            if cell.tag == "th":
                is_header = True
            cells.append(cell)
        if is_header and header_row is None:
            header_row = cells
        else:
            rows_data.append(cells)

    if not header_row and not rows_data:
        return

    ncols = len(header_row) if header_row else len(rows_data[0])
    nrows = (1 if header_row else 0) + len(rows_data)

    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = "Table Grid"

    row_idx = 0
    # Header row
    if header_row:
        for col_idx, cell_el in enumerate(header_row):
            if col_idx >= ncols:
                break
            cell = table.cell(0, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.style = doc.styles["Normal"]
            text = get_element_text(cell_el).strip()
            run = p.add_run(text)
            set_run_font(run, size=Pt(9), bold=True)
        row_idx = 1

    # Data rows
    for data_row in rows_data:
        for col_idx, cell_el in enumerate(data_row):
            if col_idx >= ncols:
                break
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.style = doc.styles["Normal"]
            render_inline(p, cell_el, font_size=Pt(9))
        row_idx += 1

    return table


def add_list(doc, list_el, ordered=False):
    """Convert an HTML <ul> or <ol> into Word paragraphs with prefixes."""
    for idx, li in enumerate(list_el.iter("li"), start=1):
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        # Add bullet or number prefix
        prefix = f"{idx}. " if ordered else "\u2022 "
        run = p.add_run(prefix)
        set_run_font(run, size=Pt(11))
        # Render the rest of the list item content
        render_inline(p, li, font_size=Pt(11))


# ── Main logic ───────────────────────────────────────────────────────────

def remove_colophon(doc):
    """Remove the trailing colophon paragraph and any preceding empty paragraphs.
    Returns the colophon paragraph's XML element (for re-adding)."""
    paragraphs = doc.paragraphs
    if not paragraphs:
        return None

    # The colophon is the last paragraph with actual text
    colophon_p = None
    colophon_idx = None
    for i in range(len(paragraphs) - 1, -1, -1):
        text = paragraphs[i].text.strip()
        if text and "Digital Critical Edition" in text:
            colophon_p = paragraphs[i]
            colophon_idx = i
            break

    if colophon_p is None:
        return None

    # Save a copy of the colophon XML
    colophon_xml = copy.deepcopy(colophon_p._element)

    # Remove colophon and any trailing empty paragraphs
    elements_to_remove = []
    for i in range(colophon_idx, len(paragraphs)):
        elements_to_remove.append(paragraphs[i]._element)

    body = doc.element.body
    for elem in elements_to_remove:
        body.remove(elem)

    # Also remove any trailing empty paragraphs before the colophon
    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        body.remove(doc.paragraphs[-1]._element)

    return colophon_xml


def re_add_colophon(doc, colophon_xml):
    """Re-add the colophon paragraph at the end of the document."""
    if colophon_xml is None:
        return
    # Add an empty paragraph for spacing
    spacer = doc.add_paragraph()
    spacer.style = doc.styles["Normal"]
    # Add the colophon
    doc.element.body.append(colophon_xml)


def process_elements(doc, elements):
    """Process the extracted HTML elements and append them to the DOCX."""
    for el in elements:
        tag = el.tag

        if tag == "h2":
            text = get_element_text(el).strip()
            add_heading_1(doc, text)

        elif tag == "h3":
            text = get_element_text(el).strip()
            add_heading_2(doc, text)

        elif tag == "p":
            # Check for special styling (summary paragraph with smaller font)
            style = el.get("style", "")
            if "font-size:0.85em" in style or "font-size:0.82em" in style:
                p = add_body_paragraph(doc, el, font_size=Pt(9))
            else:
                p = add_body_paragraph(doc, el)

        elif tag == "table":
            add_table(doc, el)

        elif tag == "ul":
            add_list(doc, el, ordered=False)

        elif tag == "ol":
            add_list(doc, el, ordered=True)

        # Skip script, style, hr, div etc.


def main():
    print("Loading HTML...")
    tree = load_html()

    print("Extracting sections VII-IX...")
    elements = find_sections_vii_to_ix(tree)
    print(f"  Found {len(elements)} elements to process")

    # Describe what we found
    tags = {}
    for el in elements:
        tags[el.tag] = tags.get(el.tag, 0) + 1
    for tag, count in sorted(tags.items()):
        print(f"    {tag}: {count}")

    print(f"Creating backup at {BACKUP_PATH}...")
    shutil.copy2(DOCX_PATH, BACKUP_PATH)

    print("Opening DOCX...")
    doc = Document(DOCX_PATH)

    print("Removing colophon...")
    colophon_xml = remove_colophon(doc)
    if colophon_xml is not None:
        print("  Colophon saved for re-insertion")
    else:
        print("  WARNING: No colophon found")

    print("Appending sections VII-IX...")
    process_elements(doc, elements)

    print("Re-adding colophon...")
    re_add_colophon(doc, colophon_xml)

    print(f"Saving to {DOCX_PATH}...")
    doc.save(DOCX_PATH)

    # Verification
    doc2 = Document(DOCX_PATH)
    headings = [p.text for p in doc2.paragraphs
                if p.style.name.startswith("Heading")]
    print(f"\nVerification: {len(headings)} headings in final document")
    for h in headings:
        print(f"  {h}")

    tables = doc2.tables
    print(f"Tables: {len(tables)}")

    # Count hyperlinks
    hyperlinks = doc2.element.body.findall(".//" + qn("w:hyperlink"))
    print(f"Hyperlinks: {len(hyperlinks)}")

    print("\nDone!")


if __name__ == "__main__":
    main()
